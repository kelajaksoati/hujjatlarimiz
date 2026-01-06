import os
import openpyxl
from openpyxl.styles import Font
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document

def smart_rename(filename):
    """Fayl nomini brendlash va tozalash"""
    base_name, extension = os.path.splitext(filename)
    # Noqonuniy belgilarni olib tashlash
    clean_name = base_name.replace("_", " ").replace("-", " ").strip()
    
    if 'bsb' in clean_name.lower():
        # BSB so'zini chiroyli formatga keltirish
        name_only = clean_name.lower().replace('bsb', '').strip()
        clean_name = f"BSB {name_only.upper()}"
    else:
        clean_name = clean_name.title()
    
    # Telegram uchun xavfsiz nom (bo'shliqlarni tagchiziq bilan almashtirish)
    final_name = f"@ISH_REJA_UZ_{clean_name.replace(' ', '_')}{extension.lower()}"
    return final_name

def edit_excel(path):
    """Excelning barcha varaqlari A1 katagiga brend yozish"""
    try:
        wb = openpyxl.load_workbook(path)
        for ws in wb.worksheets:
            # Birinchi qatorga yangi qator qo'shish
            ws.insert_rows(1)
            ws['A1'] = "@ish_reja_uz kanali uchun maxsus tayyorlandi"
            ws['A1'].font = Font(bold=True, color="FF0000", size=12)
        wb.save(path)
        wb.close()
    except Exception as e:
        print(f"Excel error: {e}")



def add_pdf_watermark(path):
    """PDF sahifalariga shaffof suv belgisi qo'shish"""
    try:
        reader = PdfReader(path)
        writer = PdfWriter()
        
        # Watermark chizish
        packet = BytesIO()
        can = canvas.Canvas(packet)
        can.setFont("Helvetica-Bold", 45)
        can.setFillGray(0.5, 0.15) # 15% shaffoflik
        can.saveState()
        
        # Sahifa markazini topish (taxminan)
        can.translate(300, 450)
        can.rotate(45)
        can.drawCentredString(0, 0, "@ish_reja_uz")
        can.restoreState()
        can.save()
        
        packet.seek(0)
        watermark_pdf = PdfReader(packet)
        watermark_page = watermark_pdf.pages[0]
        
        for page in reader.pages:
            page.merge_page(watermark_page)
            writer.add_page(page)
            
        with open(path, "wb") as f:
            writer.write(f)
    except Exception as e:
        print(f"PDF error: {e}")



def edit_docx(path):
    """Word hujjati boshiga brend matnini qo'shish"""
    try:
        doc = Document(path)
        # Hujjat bo'sh bo'lsa xato bermasligi uchun tekshirish
        if doc.paragraphs:
            doc.paragraphs[0].insert_paragraph_before("@ish_reja_uz kanali uchun maxsus tayyorlandi")
        else:
            doc.add_paragraph("@ish_reja_uz kanali uchun maxsus tayyorlandi")
        doc.save(path)
    except Exception as e:
        print(f"Docx error: {e}")
